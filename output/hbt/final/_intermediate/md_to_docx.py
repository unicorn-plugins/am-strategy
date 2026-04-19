"""Markdown → .docx 변환 (실무진 상세본용)
- 헤딩 H1~H4, 문단, 인용구, 코드블록, 리스트, GFM 표 지원
- 굵게/기울임 인라인 지원
- Calibri(본문) / 맑은 고딕(한글) 10.5pt 기본
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x59, 0x59, 0x59)
TEAL = RGBColor(0x1B, 0x6C, 0xA8)
BODY_SIZE = Pt(10.5)
BODY_FONT = "Calibri"
BODY_FONT_KR = "맑은 고딕"


def set_run_fonts(run):
    run.font.name = BODY_FONT
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rFonts.set(qn("w:eastAsia"), BODY_FONT_KR)
    rFonts.set(qn("w:cs"), BODY_FONT)


INLINE_PATTERNS = [
    (re.compile(r"\*\*(.+?)\*\*"), "bold"),
    (re.compile(r"__(.+?)__"), "bold"),
    (re.compile(r"\*(.+?)\*"), "italic"),
    (re.compile(r"`([^`]+?)`"), "code"),
    (re.compile(r"\[([^\]]+)\]\(([^)]+)\)"), "link"),
]


def add_runs(paragraph, text, base_size=BODY_SIZE, base_color=None):
    """간단한 인라인 파서: **bold**, *italic*, `code`, [text](url)"""
    idx = 0
    s = text
    while s:
        earliest = None
        earliest_kind = None
        for pat, kind in INLINE_PATTERNS:
            m = pat.search(s)
            if m and (earliest is None or m.start() < earliest.start()):
                earliest = m
                earliest_kind = kind
        if not earliest:
            run = paragraph.add_run(s)
            run.font.size = base_size
            if base_color:
                run.font.color.rgb = base_color
            set_run_fonts(run)
            break
        pre = s[:earliest.start()]
        if pre:
            run = paragraph.add_run(pre)
            run.font.size = base_size
            if base_color:
                run.font.color.rgb = base_color
            set_run_fonts(run)
        grp = earliest.group(1)
        if earliest_kind == "link":
            grp = earliest.group(1)
        run = paragraph.add_run(grp)
        run.font.size = base_size
        if base_color:
            run.font.color.rgb = base_color
        if earliest_kind == "bold":
            run.bold = True
        elif earliest_kind == "italic":
            run.italic = True
        elif earliest_kind == "code":
            run.font.name = "Consolas"
        elif earliest_kind == "link":
            run.font.color.rgb = TEAL
            run.underline = True
        set_run_fonts(run)
        s = s[earliest.end():]


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    sizes = {1: Pt(22), 2: Pt(17), 3: Pt(14), 4: Pt(12)}
    colors = {1: NAVY, 2: NAVY, 3: BLUE, 4: GRAY}
    run = p.add_run(text)
    run.bold = True
    run.font.size = sizes.get(level, Pt(12))
    run.font.color.rgb = colors.get(level, NAVY)
    set_run_fonts(run)
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.paragraph_format.page_break_before = True


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    add_runs(p, text)
    p.paragraph_format.space_after = Pt(4)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    add_runs(p, text)


def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    add_runs(p, text)


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = BODY_SIZE
    run.font.color.rgb = GRAY
    set_run_fonts(run)


def add_code(doc, code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    for i, line in enumerate(code_lines):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY


def add_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = t.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, text, base_size=Pt(9.5))
            if ri == 0:
                for r in p.runs:
                    r.bold = True
    doc.add_paragraph()


def parse_table_block(lines, start):
    """start는 헤더 라인 인덱스. 구분선(:---) 포함."""
    rows = []
    i = start
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.lstrip().startswith("|"):
            break
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if i == start + 1 and all(set(c.replace(":", "").replace(" ", "")) <= {"-"} for c in cells if c):
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


def convert(md_path: Path, out_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    raw = md_path.read_text(encoding="utf-8")
    lines = raw.split("\n")
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                add_code(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), 4)
            i += 1
            continue
        if stripped.startswith("---") or stripped.startswith("==="):
            i += 1
            continue
        if stripped.startswith(">"):
            add_quote(doc, stripped.lstrip(">").strip())
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and "-" in lines[i + 1]:
            rows, ni = parse_table_block(lines, i)
            add_table(doc, rows)
            i = ni
            continue
        m = re.match(r"^(\s*)([-*+])\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            add_bullet(doc, m.group(3), level=indent)
            i += 1
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            add_numbered(doc, m.group(2), level=indent)
            i += 1
            continue

        # 기본 문단: 다음 빈줄까지 이어붙임
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt.strip():
                break
            if nxt.lstrip().startswith(("#", ">", "|", "- ", "* ", "+ ", "```")) or re.match(r"^\s*\d+\.\s", nxt):
                break
            para_lines.append(nxt.strip())
            j += 1
        add_paragraph(doc, " ".join(para_lines))
        i = j

    doc.save(str(out_path))
    print(f"Saved: {out_path} ({out_path.stat().st_size} bytes)")


if __name__ == "__main__":
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    dst.parent.mkdir(parents=True, exist_ok=True)
    convert(src, dst)
